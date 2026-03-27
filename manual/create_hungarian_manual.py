# -*- coding: utf-8 -*-
"""
Trio Laser User Manual - angol -> magyar forditas
Minden PDF oldalt kepkent renderel, es a magyar szoveget
kulon szekcioban adja hozza minden oldal utan.
Futtatas: python create_hungarian_manual.py
"""

import fitz
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

PDF_PATH = "Laser User manual.pdf"
OUTPUT_PATH = "Trio_Laser_Felhasznaloi_Kezikonyv_HU.docx"

# Magyar forditas - oldalankent
TRANSLATIONS = {
    1: {
        "cim": "Trio Lezer Szortelenito Gep",
        "szoveg": (
            "Udvozoljuk a lezer szepsegipari szortelenito gep hasznalasaban! "
            "Ez a forradalmi dióda lezer szepségipari szortelenito keszulek sima, kényelmes, "
            "tartós hatású szortelenítési, borfiatalítási, fehérítési és érpókhalo-kezelési "
            "élményt nyújt a kulonbozo bortipusokhoz.\n\n"
            "Alkatrészek feliratai a képen:\n"
            "- Emergency Switch = Veszleallito gomb\n"
            "- 10.4' Touch Screen = Erintokepernyo\n"
            "- Key Switch = Kulcskapcsolo\n"
            "- Humanized Armrest = Ergonomikus kartamasz\n"
            "- Handle Piece = Kezelofej\n"
            "- Universal Wheel = Gorgo"
        )
    },
    2: {
        "cim": "Tartalom",
        "szoveg": (
            "Udvozlolap ............... 0\n"
            "1. Cégprofil ............. 1\n"
            "2. Figyelmeztetések ...... 2\n"
            "3. Csomaglista ........... 3\n"
            "4. Muszaki paraméterek ... 4\n"
            "5. Rendszerstruktúra ..... 5\n"
            "6. Haszanálati utasítás .. 6\n"
            "7. Kezelőfelület ......... 7\n"
            "8. Karbantartás .......... 8\n"
            "9. Garancia .............. 9\n"
            "10. Hibaelhárítás ........ 10\n"
            "11. Melléklet ............ 11\n\n"
            "1. CÉGPROFIL\n\n"
            "Az Oriental-laser (Beijing) Technology Co., Ltd. 2008-ban alapított csúcstechnológiai "
            "vállalat, amely lézerberendezések kutatás-fejlesztésére, gyártására és értékesítésére "
            "szakosodott. A székhely Pekingben, a Daxing kerületben található."
        )
    },
    3: {
        "cim": "Cégprofil (folytatás) + Figyelmeztetések",
        "szoveg": (
            "Az évek során számos szakterület kiváló szakembereit szerezte meg, elsősorban a nagy "
            "teljesítményű félvezető lézer termékek, kozmetikai berendezések fejlesztése és gyártása "
            "területén. Termékeik ipari, kozmetikai, kutatási és katonai területeken alkalmazhatók.\n\n"
            "A vállalat ISO9001 tanúsítvánnyal rendelkezik.\n\n"
            "Az integritás, szakmaiság, elkötelezettség és innováció vállalati szellem alapján "
            "az Oriental-laser kiemelkedő minőséget nyújt.\n\n"
            "2. FIGYELMEZTETÉSEK\n\n"
            "1. A gép átvételét követően 24 óra elteltével kapcsolható be: ellenőrizze a "
            "tartozékokat, szerelje össze a gépet, töltse fel vízzel, és hagyja állni 24 óráig. "
            "(Ha nem az előírásoknak megfelelően indítja el, a problémát saját maga oldja meg.)"
        )
    },
    4: {
        "cim": "Figyelmeztetések (folytatás) + Csomaglista",
        "szoveg": (
            "2. A gépbe tiszta vizet kell tölteni, amelyet havonta kell cserélni. Kb. 2,5-3 "
            "palack szükséges. Ha a készüléket több mint fél hónapig nem használják, le kell üríteni.\n\n"
            "3. Kezelés előtt figyelmesen olvassa el ezt a kézikönyvet és tanulmányozza az "
            "oktatóvideókat.\n\n"
            "4. A gép mozgatása előtt le kell üríteni. Nem dönthető meg és nem fordítható fel. "
            "Mozgatás után 24 óra elteltével indítható el.\n\n"
            "5. Az átvételkor a képernyő alján egy piros gomb látható. Forgassa el a nyíl "
            "irányába, amíg ki nem ugrik.\n\n"
            "6. Helyezze be a kulcsot és kapcsolja be a gépet.\n\n"
            "7. A páciensnek szemvédővel kell eltakarnia a szemét; a terapeutának "
            "lézervédő szemüveget kell viselnie.\n\n"
            "8. A bőr hűtése után kb. egy percet kell várni a kezelőfej lehűlésére.\n\n"
            "9. Nem kezelés közben a kezelőfejet a tartóra kell helyezni, fejjel lefelé.\n\n"
            "10. Kezelés elején és végén tisztítsa meg a kezelőfejet. Soha ne hagyjon rajta "
            "hideg gélt, szőrt vagy korpát - ez megégetheti a lézerfej ablakát.\n\n"
            "11. A gép környezetében légkondicionálónak kell lennie, a szobahőmérséklet "
            "28 C fok alatt legyen. Az állásban lévő gép 0 C fok felett legyen.\n\n"
            "3. CSOMAGLISTA\n"
            "- Szépségipari főegység: 1 db\n"
            "- Szépségipari kézdarab: 1 db\n"
            "- Lábkapcsoló: 1 db"
        )
    },
    5: {
        "cim": "Csomaglista (folytatás) + Muszaki paraméterek",
        "szoveg": (
            "Csomaglista (folytatás):\n"
            "- VAC kábel (tápkábel): 1 db\n"
            "- Vízcsoe: 1 db (vízbe- és kivezetéshez)\n"
            "- Tölcsér: 1 db\n"
            "- Kulcs: 2 db\n"
            "- Csomaglista: 1 db\n\n"
            "4. MUSZAKI PARAMÉTEREK\n\n"
            "Modell: OL-HR-ROSE\n"
            "Energia: 0-120 J/cm2\n"
            "Hullámhossz: 808nm / 755nm / 1064nm\n"
            "Impulzusszélesség: 0ms - 600ms\n"
            "Frekvenciatartomány: 1Hz - 10Hz\n"
            "Zafír kezelőfej hőmérséklete: -5 C - +5 C (hűtött zafírral)\n"
            "Lézerfej mérete: 15x25 mm\n"
            "Felület nyelve: ANGOL\n"
            "LCD képernyő mérete: 10,4 hüvelyk\n"
            "Rendszerhűtési módszer: R134A Freon kompresszoros hűtés\n"
            "Tápellátás: 220/110VAC / 50-60Hz, 10A\n"
            "Hűtővíz követelmény: Ioncserélt víz\n"
            "Környezeti hőmérséklet: 15-28 C fok\n"
            "Páratartalom: <70%\n"
            "Csomagolási méret: 123(H)x64(SZ)x64(M) cm\n"
            "Csomagolási tömeg: 65 kg\n\n"
            "5. BERENDEZÉS FELÉPÍTÉSE ÉS KAPCSOLÁSI RAJZ"
        )
    },
    6: {
        "cim": "Kapcsolási rajz + NYÁK lap leírása",
        "szoveg": (
            "A képen a rendszer kapcsolási blokkdiagramja és a vezérlőlap (NYÁK) látható.\n\n"
            "Vezérlőlap érintkezők leírása (kezdet):\n\n"
            "Szám | Leírás\n"
            "1, 2 | A vezérlőlap tápbemeneti csatlakozói. Támogatott feszültségtartomány:"
        )
    },
    7: {
        "cim": "Vezérlőlap érintkezők + Használati utasítás",
        "szoveg": (
            "9-18V. Az 1-es pozitív, a 2-es negatív.\n"
            "3, 4 | 12V-os kimenet. Alapértelmezés szerint ventilátor csatlakozik.\n"
            "5 | Kijelző csatlakozási pontja.\n"
            "6 | Számlálókártya bemeneti csatlakozója.\n"
            "7 | Frissítési program letöltési portja.\n"
            "9, 10 | 4-20mA áramsziget fogadására alkalmasak feszültségméréshez.\n"
            "11, 12 | LED-lámpák vezérlésére alkalmasak.\n"
            "13, 14, 15 | Digitális hőmérséklet-érzékelő (DS18B20) csatlakozói.\n"
            "16, 17, 18 | Hall-áramlás-érzékelő csatlakozói.\n"
            "19, 20 | Kézi kapcsolóhoz vagy pedálhoz csatlakoznak.\n"
            "21-24 | Két tartalék bemeneti kapcsolójelet biztosítanak.\n"
            "25-28 | Hűtő TEC vezérlőlábai. 27 pozitív, 28 negatív TEC táp.\n"
            "29-31 | Lézer és lézertápegység csatlakozói. Max. 100A kimenet.\n"
            "32 | Tartalék DIP-kapcsoló. Méret: 200x100x95mm\n\n"
            "6. HASZNÁLATI UTASÍTÁS\n\n"
            "Kérjük, olvassa el figyelmesen az első bekapcsolás előtt.\n\n"
            "6.1 Állási idő használat előtt\n"
            "A gép átvételét vagy mozgatását követően 24 óráig állni kell mielőtt bekapcsolja.\n\n"
            "6.1.1 Állási idő átvétel után\n"
            "Hosszú szállítás során a kompresszor kenőolaja kifolyhat. 24 óráig kell várni, "
            "hogy az olaj visszafolyjon.\n\n"
            "6.1.2 Állási idő mozgatás előtt\n"
            "Ne döntse meg és ne szállítsa hosszú távolságra. Mozgatás után 24 óráig állni kell."
        )
    },
    8: {
        "cim": "6.2 Berendezés összszerelése",
        "szoveg": (
            "6.2.1 Vízfeltöltési muvelet\n\n"
            "- Az Oriental Laser munkaállomás hűtővizét havonta kell cserélni. "
            "Csak tiszta vizet használjon (Wahaha vagy hasonló, 3 x 500 ml palack szükséges).\n\n"
            "- Túlfolyó: Illessze be a rövid kiöntőt a túlfolyó nyílásba kattanásig.\n\n"
            "- Vízbeöntő: Tartsa az előző kiöntőt a helyén; a tölcsérrel ellátott vékony végű "
            "vízcsövet illessze be a vízbeöntő nyílásba kattanásig.\n\n"
            "- Teli jelzés: Töltsön be tiszta vizet tölcsér segítségével, amíg a túlfolyóból "
            "víz nem folyik ki - ekkor a tartály tele van.\n\n"
            "Megjegyzés a képen:\n"
            "- Water Inlet = Vízbemenet\n"
            "- Water Spilt = Túlfolyó\n"
            "- Water Outlet = Vízkimenet\n"
            "A tartály kb. 3 litert tartalmaz. Legalább 2,5-3,0 liter lepárolt vizet adjon hozzá. "
            "A víz minősége legyen kristálytiszta; csak purifikált vagy lepárolt vizet használjon.\n\n"
            "6.2.2 Leürítési muvelet\n\n"
            "- Csatlakoztassa a leürítőt: illessze be a rövid kiöntőt a leürítő nyílásba kattanásig."
        )
    },
    9: {
        "cim": "6.2 Összszerelés (folytatás) + 6.3 Előkészítés",
        "szoveg": (
            "- Leürítési muvelet: A túlfolyó vízcsoe beillesztése után automatikusan elindul a "
            "vízelvezetés. Medencével fogja fel a lefolyó vizet.\n"
            "- Befejezés: Amikor nem folyik több víz, a tartály üres.\n\n"
            "6.2.3 Kézdarab-tartó felszerelése\n"
            "A berendezés oldalán három csavarlyukon keresztül, süllyesztett fejű csavarokkal "
            "szerelje fel a kézdarab-tartót.\n\n"
            "6.2.4 Kézdarab csatlakoztatása\n"
            "Ujjaival nyomja meg a kézdarab csatlakozójának két pontját és nyomja határozottan "
            "a csatlakozóba kattanásig - ellenőrizze, hogy nincs vízszivárgás.\n"
            "Leszerelésnél: nyomja meg a két pontot és húzza ki. "
            "A gyenge csatlakozás riasztást és szivattyúkárosodást okozhat.\n\n"
            "6.2.5 Tápcsatlakozó\n"
            "A tápkábel egyik végét a főegységhez, másik végét az elosztóhoz csatlakoztassa.\n"
            "Első alkalommal távolítsa el a képernyő védőfóliáját.\n\n"
            "6.3 SZORTELENÍTÉS ELŐKÉSZÍTÉSE\n\n"
            "6.3.1 Fotó készítése\n"
            "Minden kezelés előtt (borotválkozás előtt is) jó megvilágításban készítsen fényképet "
            "az összehasonlíthatóság érdekében.\n\n"
            "6.3.2 Rész borotválása\n"
            "- Az ügyfélnek egy hétig kerülnie kell a napsugárzást;\n"
            "- A szortelenítési területen lévő sminket le kell távolítani;"
        )
    },
    10: {
        "cim": "6.3-6.4 Szortelenítési muveletek",
        "szoveg": (
            "- Utasítsa az ügyfelet a szőr óvatos borotválására - ne használjanak gyantát, "
            "epilátort, vagy szortelenítő krémet.\n\n"
            "6.3.3 Bőrhűtés\n"
            "Vigyen fel vastag réteg hideg gélt a szortelenítési területre. Érzékeny "
            "területeknél kezelés előtt jégcsomagot is alkalmazhat.\n\n"
            "6.4 LÉZER SZORTELENÍTÉSI KEZELÉS\n\n"
            "6.4.1 Bekapcsolás\n"
            "(1) Fordítsa a kulcsot jobbra a bekapcsoláshoz. Vészhelyzet esetén nyomja meg "
            "a piros vészleállító gombot. Ha korábban megnyomta, fél fordulattal jobbra kell "
            "visszaállítani az újraindításhoz.\n"
            "(2) Muködés közben a vészleállítóval azonnali leállítás lehetséges.\n"
            "(3) Kétnyelvű felületen válasszon nyelvet az udvözlőképernyőn.\n"
            "(4) Kikapcsoláshoz forgassa vissza a kulcsot, biztonság érdekében húzza ki a "
            "tápkábelt.\n\n"
            "6.4.2 Kezelési mód kiválasztása\n"
            "Három üzemmód:\n"
            "- HR (Hair Removal) = Szortelenítés\n"
            "- FHR (Fast Hair Removal) = Gyors szortelenítés\n"
            "- SR (Skin Rejuvenation) = Bőrfiatalítás\n\n"
            "Először a 'Professzionális szortelenítés' módot ismertetjük részletesen."
        )
    },
    11: {
        "cim": "6.4.3-6.4.5 Paraméterek és kezelés",
        "szoveg": (
            "6.4.3 Funkció kiválasztása\n"
            "(1) nem, (2) bőrszín, (3) szortelenítési terület.\n"
            "Az ügyfél állapota alapján válasszon, majd kattintson a Megerősítés gombra.\n\n"
            "6.4.4 Paraméterbeállítás\n"
            "(1) Sötét, sűrű szor és érzékeny területeknél: energia 6J-ről kezdve, "
            "frekvencia 1-ről.\n"
            "(2) Világos, ritka szor és érzéketlen területeknél: energia 10J-ről kezdhető, "
            "frekvencia 2-ről vagy 3-ról.\n"
            "(3) Az impulzusszélességet nem kell beállítani - automatikusan igazodik.\n"
            "(4) A bőrhűtési paramétert 80-90%-ra kell növelni.\n\n"
            "6.4.5 Kezelési muvelet\n"
            "A beállítások elvégzése után kattintson a 'Kesz' gombra. Az előkészítési idő "
            "legfeljebb 30 másodperc; ha tovább tart, a rendszer készenléti állapotba vált.\n\n"
            "(1) Kezelés közben folyamatosan igazítsa az energiát és a frekvenciát. "
            "Javasolt az ügyfél által elviselhető maximális energiasűrűséggel csúsztató "
            "technikával dolgozni. A kezelőfej teljes felületen érintkezzen a bőrrel. "
            "Az optimális hatás: enyhe szúrás vagy hangyacsípés érzése."
        )
    },
    12: {
        "cim": "6.4.5-6.4.8 Kezelés folytatása",
        "szoveg": (
            "(2) Jegyezze fel a beállított lézerparamétereket - a következő kezelésnél "
            "ugyanezeket lehet alkalmazni.\n"
            "(3) Minél magasabb az energia és frekvencia, annál jobb az eredmény. Az ügyfél "
            "által elviselhető enyhe szúrás az optimális - túl magas energia égési sérülést okoz.\n"
            "(4) Általánosságban csúsztató módszerrel dolgozzon.\n"
            "(5) Az egyes testrészekre vonatkozó értékek a paramétertáblázatban találhatók.\n\n"
            "6.4.6 Gyors szortelenítés\n"
            "(1) Csak a bőrhűtést kell beállítani (80-90%-ra). A többi paramétert a gép "
            "előre beállítja.\n"
            "(2) Minden testterületnél a legalacsonyabb energiaszintről kell kezdeni.\n\n"
            "6.4.8 A lézer szortelenítés elve\n"
            "(1) A lézer szortelenítés a melanin fényelnyelő tulajdonságát használja fel: "
            "a melanin absorbeálja a lézerfényt, hőhatást generál, és elpusztítja a "
            "szortüszőt. Mivel a szortüsző pusztulása visszafordíthatatlan, a lézer "
            "tartós szortelenítést biztosít.\n\n"
            "(2) A szornek növekedési ciklusa van: növekedés - hanyatlás - nyugalmi szakasz. "
            "A hanyatlás és nyugalmi szakaszban lévő szortüszők nem tartalmaznak melanint. "
            "Javasolt 21-30 napos kezelési időköz 5-8 alkalmon keresztül, több mint fél éven "
            "át. Ha marad finom szor, 2-3 megerősítő kezeléssel elérhető a teljes tartós "
            "szortelenítés."
        )
    },
    13: {
        "cim": "6.4.9 Utókezelés + 6.5 Bőrfiatalítás",
        "szoveg": (
            "6.4.9 Kezelés utáni gondozás\n"
            "(1) Szortelenítés után jégcsomagot vagy aloe vera gélt alkalmazhat.\n"
            "(2) Ne változtasson a napi bőrápolási rutinján; kerülje az irritáló kozmetikumokat.\n"
            "(3) Ha a bőrben implantátum van, vagy mikroinjekciós kezelést végeztek, a lézer "
            "kezelés nem ajánlott.\n"
            "(4) Néhány ügyfélen kiütés jelentkezhet - ez normális, 3-7 nap alatt elmúlik.\n"
            "(5) A kezelt területet 1 napig kerüljük a forró vízzel, 1 hétig a napsugárzással.\n"
            "(6) Alkohol fogyasztása és csípős, irritáló ételek kerülendők.\n"
            "(7) Gyantás szortelenítés után lézeres kezelés csak 6 hónap után alkalmazható.\n"
            "(8) Az 1-2. kezelés után a szor kihúzható lehet - de ne tegye! A szor nélküli "
            "szortüsző nem vezet hőt, ami rontja a szortelenítés hatékonyságát.\n\n"
            "6.5 LÉZER BŐRFIATALÍTÁSI KEZELÉS\n\n"
            "6.5.1 Kezelési módszer\n"
            "(1) Csak a bőrhűtést kell beállítani (80-90%-ra).\n"
            "(2) Új ügyfelek esetén a 'Lagy fehérítés' üzemmóddal kezdjen.\n\n"
            "6.5.2 Elv\n"
            "Az Oriental-Laser fő hullámhossza 808-810nm. A melanin, hemoglobin és víz "
            "absorbeálja ezt a közeli infravörös fényt. A hőhatás eredményei:\n"
            "(1) A lézer új kollagén sejtek szintézisét indukálja a dermiszben, ráncokat "
            "és hegesedéseket javít."
        )
    },
    14: {
        "cim": "6.5 Bőrfiatalítás (folytatás) + 7. Kezelőfelület paraméterei",
        "szoveg": (
            "(2) A lézer melanint és kapillárisokat céloz meg hőbontás céljából. Csökkenti "
            "a pigmentációt, mérsékeli a telangiektáziát, javítja a bőr simaságát. "
            "Kezelhető problémák: fotóöregedés, öregségi foltok, érpókháló, szürke bőr, "
            "rosacea.\n\n"
            "6.5.3 Utókezelés\n"
            "(1) Teljes arckezelés ideje: 10-20 perc (enyhe állapot: 10 perc, súlyosabb: "
            "20 perc). Két kezelés között általában 15 nap.\n"
            "(2) A betegek több mint 90%-ánál szignifikáns javulás tapasztalható.\n"
            "(3) Nincs leállási idő. Ritkán ödéma vagy átmeneti pigmentációs változás "
            "fordulhat elő, ami magától visszatér.\n"
            "(4) A kezelés után új kollagén szintézis mutatható ki.\n\n"
            "7. KEZELŐFELÜLET PARAMÉTEREINEK MEGHATÁROZÁSA\n\n"
            "7.1 Fluence (Energiasűrűség): Az egységnyi területre eső egyetlen lézerimpulzus "
            "energiaértéke. 1-120J között állítható.\n\n"
            "7.2 Frekvencia: Másodpercenkénti lézerimpulzusok száma. 1Hz-10Hz között "
            "állítható 1Hz-es lépésközzel."
        )
    },
    15: {
        "cim": "7. Kezelőfelület paraméterei (folytatás)",
        "szoveg": (
            "A frekvencia a '+' és '-' gombokkal is állítható.\n\n"
            "7.3 Részösszeg: Az LCD képernyő jobb felső sarkában a kezelés teljes "
            "lézerenergiáját mutatja. A 'Torles' gombbal nullázható.\n\n"
            "7.4 Bőrhűtési szint: A kezelőfej zafír ablakának és alumínium végének hűtési "
            "hatékonyságát jelzi. 0-5 C fokos bőrfelületi hőmérsékletet tart fenn. "
            "'+' és '-' gombokkal állítható. Minél hosszabb a zöld csík, annál erősebb a hűtés. "
            "Bekapcsoláskor alapértelmezés szerint ki van kapcsolva.\n\n"
            "7.5 Állapot (STATUS):\n"
            "7.5.1 STANDBY (Készenlét) - Ilyenkor állítható a bőrhűtési szint. "
            "Kesz vagy Mukodik állapotban nem.\n"
            "7.5.2 READY (Kesz) - A 'READY' gomb megnyomása után. A lézer készen áll; "
            "a kézdarab gombja vagy a lábkapcsoló aktiválja.\n"
            "7.5.3 WORKING (Mukodik) - A kézdarab gombja vagy a lábkapcsoló aktív.\n"
            "7.5.4 Az állapotsor riasztásjelzéseket is megjelenít meghibásodás esetén.\n\n"
            "7.6 Képernyő alján megjelenő adatok:\n"
            "7.6.1 Vízhőmérséklet (WATER TEMP): A hűtővíz hőmérséklete.\n"
            "7.6.2 Áramlás (FLOW RATE): A hűtővíz áramlási sebessége.\n"
            "7.6.3 Ionitás (CONDUCTIVITY): A hűtővíz vezetőképessége és tisztasága."
        )
    },
    16: {
        "cim": "7.6.4 + 8. Karbantartás + 9. Garancia + 10. Hibaelhárítás",
        "szoveg": (
            "7.6.4 Összes lövés (TOTAL): A kezelőfej összes lövésének kumulatív száma.\n\n"
            "8. KARBANTARTÁS\n\n"
            "8.1 Csak tiszta vizet töltsön be. Ellenkező esetben a készülék gyorsan megrongálódhat.\n\n"
            "8.2 A gépet légkondicionált helyiségben kell üzemeltetni (5-28 C fok).\n"
            "1. A hőmérséklet ne csökkenjen nulla alá - jég képződhet.\n"
            "2. Tisztítsa meg a gélt a kézdarabról; ne legyen por a zafír körül.\n"
            "3. Ne hajlítsa meg erősen a kézdarab csövét.\n"
            "4. Ne csatlakoztassa/válassza le a kézdarabot túl gyakran.\n"
            "5. Ne süsse el véletlenszerűen a lézert - égési sérülést vagy tüzet okozhat.\n"
            "6. Mindig tartsa a kézdarabot a tartón, ha nem használja.\n\n"
            "9. GARANCIÁLIS FELTÉTELEK\n\n"
            "Nem érvényes a garancia az alábbi esetekben:\n"
            "1. Víz nélküli üzemeltetés.\n"
            "2. Bekapcsolás a kézdarab csatlakoztatása nélkül.\n"
            "3. A kézdarab rendszeres leejtése.\n"
            "4. Helytelen bemeneti feszültség.\n"
            "5. Egyéb kezelési hibák.\n\n"
            "10. GYORS HIBAELHÁRÍTÁS\n\n"
            "A bekapcsolás után a rendszer folyamatosan önellenőrző programot futtat. "
            "Ha bármely összetevőben rendellenesség tapasztalható, riasztás jelenik meg:"
        )
    },
    17: {
        "cim": "10.1-10.2 Riasztások",
        "szoveg": (
            "10.1 AZ INDÍTÁSI ÜDVÖZLŐKÉPERNYŐN MEGJELENŐ RIASZTÁS:\n\n"
            "Riasztás: A lézerkézdarab nincs megfelelően beillesztve.\n"
            "Értelmezés: A kézdarab nincs megfelelően beillesztve, rossz a kontaktus.\n"
            "Megoldás:\n"
            "1. Ellenőrizze a kézdarab csatlakozását.\n"
            "2. Ha a dugaszon belüli kis áramköri lap megrongálódott, vegye fel a kapcsolatot "
            "a gyártóval.\n"
            "Program frissítése a riasztás maszkolásához (vészhelyzeti megoldás, nem ajánlott).\n\n"
            "10.2 A BEÁLLÍTÁSI FELÜLETEN MEGJELENŐ RIASZTÁSOK:\n\n"
            "Riasztás: Magas vízhőmérséklet / Túlhőmérséklet-riasztás.\n"
            "A hűtővíz hőmérsékletének 15-28 C fok között kell lennie. Ha meghaladja a "
            "34 C fokot, a lézer leáll.\n"
            "Megoldás:\n"
            "1. Ha a kompresszor rendesen muködik, de a hőmérséklet magas: csökkentse a "
            "szobahőmérsékletet 30 C fok alá, és várjon kb. fél órát.\n"
            "2. Ellenőrizze a kompresszor muködését (érintse meg, érez-e rezgést).\n"
            "3. Ellenőrizze, hogy elegendő víz van-e a tartályban. Ha kevés, a kompresszor "
            "relé kibillenhet."
        )
    },
    18: {
        "cim": "10.2 Riasztások (folytatás)",
        "szoveg": (
            "Ha elegendő víz van és a kompresszor normálisan muodik, de a hőmérő 34 fok "
            "alatt mutat, a hőérzékelő vagy a vezérlőlap megrongálódhatott.\n"
            "1) Ellenőrizze a szonda és az áramkör közötti kapcsolatot.\n"
            "2) Cserélje ki a lapot.\n"
            "3) Ha a csere sem oldja meg, a szonda megrongálódott.\n\n"
            "Riasztás: Alacsony vízáramlás\n"
            "Ha az interfész '0 L/perc' értéket mutat:\n"
            "1) A készülék kézdarab nélkül indult el - illessze be.\n"
            "2) Laza kapcsolat az áramlásérzékelő és a vezérlőlap között.\n"
            "3) A lap megrongálódott (általában az OW1 optocoupler).\n"
            "4) Az áramlásérzékelő megrongálódott.\n\n"
            "Ha 1,6 L/perc alatt mutat:\n"
            "1) A kézdarab nincs megfelelően érintkezve a csatlakozóval - csatlakoztassa újra.\n"
            "2) A vízkeringési rendszer légnyomása egyensúlyhiányban van - "
            "nyissa ki a vízbeöntő csatlakozót a kiegyenlítéshez.\n\n"
            "Riasztás: Kérjük, cserélje ki a hűtővizet és az ionszűrőt.\n"
            "1. Ha az ionmérő 20 felett mutat: cserélje a vizet tiszta vízre.\n"
            "2. Ha 20 alatt mutat, de a riasztás megmarad: ellenőrizze a csatlakozást "
            "vagy frissítse a programot.\n\n"
            "Riasztás: 2-es hiba (ERROR 02)\n"
            "1. Ellenőrizze a vezérlőlap hűtőborda ventilátorát.\n"
            "2. Gyenge hőelvezetés vagy túl magas hőmérséklet - lehűlés után indítsa újra.\n\n"
            "10.3 A 'KESZ' GOMB MEGNYOMÁSA UTÁN MEGJELENŐ RIASZTÁS:"
        )
    },
    19: {
        "cim": "10.3-10.4 Riasztások",
        "szoveg": (
            "Riasztás: A lézermodul megszakadt.\n"
            "Értelmezés: A lézermodul és a kézdarabban lévő kábel közötti kapcsolat megszakadt.\n"
            "Megoldás:\n"
            "1. Ellenőrizze a kábelt a vezérlőlaptól a lézermodulig.\n"
            "2. Ha a lézermodul belső megszakadás van, cserélje ki. Vegye fel a kapcsolatot "
            "a gyártóval.\n\n"
            "10.4 A KÉZDARAB GOMBJÁNAK MEGNYOMÁSAKOR MEGJELENŐ RIASZTÁS:\n\n"
            "Riasztás: Helytelen mukovedes.\n"
            "A lézer aktiválása előtt 'Kesz' állapotban kell lenni.\n"
            "Megoldás: Engedje fel a kapcsolót, és nyomja meg a képernyőn a 'Kesz' gombot.\n\n"
            "Riasztás: Kérjük, ellenőrizze a dióda lézert vagy a dióda meghajtót.\n"
            "A rendszer azt észlelte, hogy 2 lézerrúdnál több megrongálódott, vagy a "
            "kapcsolóüzemű tápegység kimeneti feszültsége túl magas.\n"
            "Megoldás:\n"
            "1. Kapcsolja ki, indítsa újra, nézze meg hány lézerrúd világít halványan.\n"
            "2. Ha 3 vagy több megrongálódott, vegye fel a kapcsolatot a gyárral.\n"
            "3. Ha csak 1-2 nem világít, frissítse az alaplaprogramot, vagy csökkentse "
            "a tápegység kimeneti feszültségét 2V-tal minden megrongálódott rúd után "
            "(vészhelyzeti megoldás)."
        )
    },
    20: {
        "cim": "10.4 Riasztások (folytatás)",
        "szoveg": (
            "Riasztás: Kérjük, ellenőrizze a dióda lézert (folytatás)\n\n"
            "Ez vészhelyzeti megoldás - hosszú távú használatra új lézer cseréje ajánlott.\n\n"
            "[A táblázat az eredeti angol képen látható a fenti képen.]"
        )
    },
    21: {
        "cim": "Vége / Melléklet",
        "szoveg": (
            "Ez vészhelyzeti megoldás - hosszú távú használatra új lézer cseréje ajánlott.\n\n"
            "A dokumentum vége."
        )
    },
}


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_p(doc, text):
    doc.add_paragraph(text)


def render_page(pdf_doc, page_num, dpi=110):
    page = pdf_doc[page_num]
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat)
    return pix.tobytes("png")


def build():
    doc = Document()

    # Margók
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    pdf_doc = fitz.open(PDF_PATH)
    total = len(pdf_doc)
    print(f"PDF oldalak szama: {total}")

    for i in range(total):
        page_num = i + 1
        print(f"  Oldal {page_num}/{total}...")
        tr = TRANSLATIONS.get(page_num, {})

        # --- Magyar szöveg ---
        title = tr.get("cim", f"{page_num}. oldal")
        add_h(doc, f"{page_num}. oldal — {title}", level=2)

        szoveg = tr.get("szoveg", "")
        if szoveg:
            add_p(doc, szoveg)

        # --- Eredeti PDF oldal képként ---
        img_bytes = render_page(pdf_doc, i, dpi=110)
        img_stream = io.BytesIO(img_bytes)
        doc.add_picture(img_stream, width=Inches(5.8))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Oldaltörés (utolsó oldal kivételével)
        if i < total - 1:
            doc.add_page_break()

    pdf_doc.close()
    doc.save(OUTPUT_PATH)
    print(f"\nKesz! Mentve: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()

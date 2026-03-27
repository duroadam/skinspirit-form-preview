# -*- coding: utf-8 -*-
"""
Trio Laser - Felhasználói Kézikönyv (Magyar)
Professzionális Word dokumentum generátor
Futtatás: python build_manual_hu.py
"""

import fitz
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os

PDF_PATH  = "Laser User manual.pdf"
IMG_DIR   = "extracted_images"
OUT_PATH  = "Trio_Laser_Kezikonyv_Magyar.docx"

# ─────────────────────────────────────────────────────────────
# Segédfüggvények
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def bold_run(para, text, size=11, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_body(doc, text, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_img(doc, path, width_inch=5.0, center=True):
    if not os.path.exists(path):
        doc.add_paragraph(f"[Kép nem található: {path}]")
        return
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inch))
    return p

def add_img_bytes(doc, img_bytes, width_inch=5.0, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Inches(width_inch))
    return p

def render_page_region(pdf_path, page_num, clip_rect=None, dpi=130):
    """Renderel egy PDF oldalt (vagy annak egy részét) PNG byte-ként."""
    doc = fitz.open(pdf_path)
    page = doc[page_num]
    mat  = fitz.Matrix(dpi / 72, dpi / 72)
    if clip_rect:
        clip = fitz.Rect(*clip_rect)
        pix  = page.get_pixmap(matrix=mat, clip=clip)
    else:
        pix  = page.get_pixmap(matrix=mat)
    doc.close()
    return pix.tobytes("png")

def img(name):
    return os.path.join(IMG_DIR, name)

def add_header_line(doc, text="Trio Lézer Szőrtelenítő"):
    """Fejléc vonal minden oldal tetején."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.underline = True
    return p

def two_col_table(doc, rows, col_widths=(2.2, 4.1)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        t.rows[i].cells[0].width = Inches(col_widths[0])
        t.rows[i].cells[1].width = Inches(col_widths[1])
        for c in t.rows[i].cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    return t

def three_col_table(doc, header, rows, col_widths=(1.1, 2.5, 2.7)):
    t = doc.add_table(rows=1+len(rows), cols=3)
    t.style = "Table Grid"
    # fejléc
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "2C3E50")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # sorok
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = str(val)
            for para in t.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    return t

# ─────────────────────────────────────────────────────────────
# FŐ DOKUMENTUM FELÉPÍTÉSE
# ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Oldalméret: A4
    for sec in doc.sections:
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ══════════════════════════════════════
    # BORÍTÓ
    # ══════════════════════════════════════
    add_header_line(doc)
    add_img(doc, img("p1_img0.jpeg"), width_inch=4.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bold_run(t, "Trio Lézer Szőrtelenítő Gép", size=20)

    doc.add_paragraph()
    welcome = doc.add_paragraph()
    welcome.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = welcome.add_run(
        "Üdvözöljük a lézer szépségipari szőrtelenítő gép használatában! "
        "Ez a forradalmi dióda lézer szépségipari szőrtelenítő készülék sima, kényelmes, "
        "tartós hatású szőrtelenítési, bőrfiatalítási, fehérítési és érpókháló-kezelési "
        "élményt nyújt a különböző bőrtípusokhoz."
    )
    r.font.size = Pt(11)

    # Alkatrész-feliratok magyarázata
    doc.add_paragraph()
    p = doc.add_paragraph()
    bold_run(p, "A főegység alkatrészei (a képen jelölve):", size=10)
    for en, hu in [
        ("Emergency Switch", "Vészleállító gomb"),
        ("10.4\u2019 Touch Screen", "10,4\u201d Érintőképernyő"),
        ("Key Switch", "Kulcskapcsoló"),
        ("Humanized Armrest", "Ergonomikus kartámasz"),
        ("Handle Piece", "Kezelőfej"),
        ("Universal Wheel", "Görgő (mozgatókerék)"),
    ]:
        add_bullet(doc, f"{en}  →  {hu}", size=10)

    doc.add_page_break()

    # ══════════════════════════════════════
    # TARTALOM
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "Tartalom", level=1)
    contents = [
        ("Üdvözlőlap", "0"),
        ("1. Cégprofil", "1"),
        ("2. Figyelmeztetések", "2"),
        ("3. Csomaglista", "3"),
        ("4. Műszaki paraméterek", "4"),
        ("5. Berendezés felépítése és kapcsolási rajz", "5"),
        ("6. Használati utasítás", "6"),
        ("7. Kezelőfelület paramétereinek meghatározása", "7"),
        ("8. Karbantartás", "8"),
        ("9. Garanciális feltételek", "9"),
        ("10. Egyszerű hibaelhárítás", "10"),
        ("11. Melléklet", "11"),
    ]
    for title, page in contents:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(title)
        r1.font.size = Pt(11)
        # pont-kitöltés és oldalszám
        r2 = p.add_run(" " + "." * (55 - len(title)) + f" {page}")
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 1. CÉGPROFIL
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "1. Cégprofil", level=1)
    add_body(doc,
        "Az Oriental-laser (Beijing) Technology Co., Ltd. 2008-ban alapított csúcstechnológiai "
        "vállalat, amely lézerberendezések kutatás-fejlesztésére, gyártására és értékesítésére "
        "szakosodott. A székhelye és K+F központja Pekingben, a Daxing kerületben - Kína "
        '"Szil\u00edcium-v\u00f6lgy\u00e9ben" - tal\u00e1lhat\u00f3, \u00e9s hongkongi fi\u00f3kirod\u00e1val rendelkezik.'
    )
    add_body(doc,
        "Az évek során számos szakterület kiváló szakembereit szerezte meg, elsősorban a nagy "
        "teljesítményű félvezető oldalpumpa modulok (DPSS lézer), nagy teljesítményű "
        "félvezető lézer termékek, félvezető lézer kozmetikai berendezések fejlesztése és "
        "gyártása, valamint rendszerintegráció területén. Termékeik ipari, kozmetikai, "
        "kutatási és katonai területeken egyaránt széles körben alkalmazhatók."
    )
    add_body(doc,
        "A vállalat 500 négyzetméteres, összességében 10K osztályú, helyileg 1000 osztályú "
        "ultratiszta műhellyel rendelkezik, amely szigorúan megfelel az US 109E szövetségi "
        "ipari szabványoknak, és ISO9001 minőségirányítási tanúsítvánnyal rendelkezik."
    )
    add_body(doc,
        'Az "integritás, szakmaiság, elkötelezettség és innováció" vállalati szellem alapján '
        "az Oriental-laser tudományos menedzsmenttel, élvonalbeli technológiával, "
        "szabványosított gyártással és kiváló szolgáltatással kíván vezető szerepet betölteni. "
        "A vállalat nagy hangsúlyt fektet a termékinnováicóra, a folyamatellenőrzésre és a "
        "minőségirányításra."
    )

    doc.add_page_break()

    # ══════════════════════════════════════
    # 2. FIGYELMEZTETÉSEK
    # ══════════════════════════════════════
    add_header_line(doc)
    add_img(doc, img("p3_img1.jpeg"), width_inch=0.8)

    add_heading(doc, "2. Figyelmeztetések", level=1)
    warnings = [
        ("1.", "A gép átvételét követően 24 óra elteltével kapcsolható be: ellenőrizze a "
               "tartozékokat, szerelje össze a gépet, töltse fel vízzel, és hagyja állni "
               "24 óráig. (Ha nem az előírásoknak megfelelően indítja el, a problémát saját "
               "maga oldja meg.)"),
        ("2.", "A gépbe tiszta vizet kell tölteni, amelyet havonta kell cserélni. Kb. 2,5-3 "
               "palack szükséges. Ha a készüléket több mint fél hónapig nem használják, le "
               "kell üríteni."),
        ("3.", "A gép átvétele előtt figyelmesen olvassa el a kezelési eljárást, tanulmányozza "
               "az oktatóvideókat és az egyéb információkat, majd kapcsolja be."),
        ("4.", "A gép mozgatása előtt le kell üríteni. Nem dönthető meg és nem fordítható fel. "
               "Mozgatás után 24 óra elteltével indítható el. Hosszú távú szállítás előtt "
               "értesítse az értékesítési munkatársat."),
        ("5.", "Az átvételkor a képernyő alján egy piros gomb látható. Forgassa el a gombon "
               "lévő nyíl irányába, amíg ki nem ugrik."),
        ("6.", "Helyezze be a kulcsot és kapcsolja be a gépet."),
        ("7.", "A páciensnek szemvédővel kell teljesen eltakarnia a szemét; a terapeutának "
               "lézervédő szemüveget kell viselnie."),
        ("8.", "A bőr hűtése után kb. egy percet kell várni a kezelőfej lehűlésére."),
        ("9.", "Nem kezelés közben a kezelőfejet a főegység oldalán lévő tartóra kell "
               "helyezni, fejjel lefelé - hogy a kondenzvíz ne folyjon vissza a kézdarabba."),
        ("10.", "Kezelés elején és végén tisztítsa meg a kezelőfejet. Soha ne hagyjon rajta "
                "hideg gélt, szőrt vagy korpát - ez megégetheti a lézerfej ablakát."),
        ("11.", "A gép környezetében légkondicionálónak kell lennie, a szobahőmérséklet "
                "28 °C alatt legyen. Az állásban lévő gép 0 °C fölött legyen - ha a belső "
                "vízkeringési rendszer befagy, az közvetlen eszközkárosodást okoz."),
    ]
    t = doc.add_table(rows=len(warnings), cols=2)
    t.style = "Table Grid"
    # nincs keret az első oszlopon
    for i, (num, text) in enumerate(warnings):
        c0 = t.rows[i].cells[0]
        c1 = t.rows[i].cells[1]
        c0.width = Inches(0.4)
        c1.width = Inches(5.9)
        c0.text = num
        c1.text = text
        for cell in (c0, c1):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        # bold a sorszámra
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════
    # 3. CSOMAGLISTA
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "3. Csomaglista", level=1)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    fejlec = ["Terméknév", "Mennyiség", "Megjegyzés"]
    for j, h in enumerate(fejlec):
        t.rows[0].cells[j].text = h
        set_cell_bg(t.rows[0].cells[j], "2C3E50")
        for para in t.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    csomaglista = [
        ("Szépségipari főegység", "1 db", ""),
        ("Szépségipari kézdarab (kezelőfej)", "1 db", ""),
        ("Lábkapcsoló", "1 db", ""),
        ("VAC kábel (tápkábel)", "1 db", ""),
        ("Vízcsövek", "1 db", "vízbe- és kivezetéshez"),
        ("Tölcsér", "1 db", ""),
        ("Kulcs", "2 db", ""),
        ("Csomaglista", "1 db", ""),
    ]
    for sor in csomaglista:
        row = t.add_row()
        for j, val in enumerate(sor):
            row.cells[j].text = val
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 4. MŰSZAKI PARAMÉTEREK
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "4. Műszaki paraméterek", level=1)
    add_heading(doc, "4.1 Elektromos paraméterek", level=2)

    params = [
        ("Modell",                           "OL-HR-ROSE"),
        ("Energia",                          "0-120 J/cm²"),
        ("Hullámhossz",                      "808 nm / 755 nm / 1064 nm"),
        ("Impulzusszélesség tartomány",      "0 ms - 600 ms"),
        ("Frekvenciatartomány",              "1 Hz - 10 Hz"),
        ("Zafír kezelőfej hőmérséklete",     "−5 °C - +5 °C (hűtött zafírral)"),
        ("Lézerfej mérete",                  "15 × 25 mm"),
        ("Felület nyelve",                   "ANGOL"),
        ("LCD képernyő mérete",              "10,4\""),
        ("Rendszerhűtési módszer",           "R134A Freon kompresszoros hűtés"),
        ("Tápellátás",                       "220/110 VAC / 50-60 Hz, 10 A"),
        ("Hűtővíz követelmény",              "Ioncserélt (desztillált) víz"),
        ("Környezeti hőmérséklet",           "15-28 °C"),
        ("Páratartalom",                     "<70 %"),
        ("Csomagolási méret",                "123 (H) × 64 (Sz) × 64 (M) cm"),
        ("Csomagolási tömeg",                "65 kg"),
    ]
    two_col_table(doc, params)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 5. BERENDEZÉS FELÉPÍTÉSE ÉS KAPCSOLÁSI RAJZ
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "5. Berendezés felépítése és kapcsolási rajz", level=1)

    # A kapcsolási blokk diagram (a PDF 6. oldalán vektoros - rendereljük)
    diagram_bytes = render_page_region(PDF_PATH, 5, clip_rect=(40, 30, 540, 250), dpi=140)
    add_img_bytes(doc, diagram_bytes, width_inch=5.5)
    add_body(doc, "A fenti ábra a rendszer kapcsolási blokkdiagramját mutatja.", size=9)

    add_img(doc, img("p6_img2.jpeg"), width_inch=5.5)  # PCB board foto
    add_body(doc, "A vezérlőlap (NYÁK) elrendezése és csatlakoztatási pontjai.", size=9)

    # Tüskék leírása
    add_heading(doc, "Vezérlőlap érintkezők leírása", level=2)
    tusk = [
        ("Sorszám", "Leírás"),
        ("1, 2",    "Tápbemeneti csatlakozók. Feszültségtartomány: 9-18 V. Az 1-es (+), 2-es (−)."),
        ("3, 4",    "12 V-os kimenet. Alapértelmezés szerint a ventilátor csatlakozik ide."),
        ("5",       "Kijelző csatlakozási pontja."),
        ("6",       "Számlálókártya bemeneti csatlakozója."),
        ("7",       "Frissítési program letöltési portja."),
        ("9, 10",   "4-20 mA áramsziget fogadására alkalmasak feszültségméréshez."),
        ("11, 12",  "LED-lámpák vezérlésére alkalmasak."),
        ("13, 14, 15", "Digitális hőmérséklet-érzékelő (DS18B20) csatlakozói."),
        ("16, 17, 18", "Hall-elven működő áramlás-érzékelő csatlakozói."),
        ("19, 20",  "Kézi kapcsolóhoz vagy lábkapcsolóhoz csatlakoznak."),
        ("21-24",   "Két tartalék bemeneti kapcsolójelet biztosítanak (21-22. és 23-24. pár)."),
        ("25-28",   "Hűtő TEC vezérlőlábak. 27-es TEC táp (+), 28-as (−). A 25. és 26. a TEC-hez vezet. "
                    "A TEC vezérlőmodul a legtöbb TEC típushoz kompatibilis; a hűtési kapacitás állítható."),
        ("29-31",   "Lézer és lézertápegység csatlakozói. 29-es: lézer (−), 30-as: lézer (+) és tápegység (+), "
                    "31-es: tápegység (−). A vezérlőlap max. 100 A kimenetet támogat."),
        ("32",      "Tartalék DIP-kapcsoló. Méret: 200 × 100 × 95 mm"),
    ]
    t2 = doc.add_table(rows=len(tusk), cols=2)
    t2.style = "Table Grid"
    for i, (k, v) in enumerate(tusk):
        t2.rows[i].cells[0].text = k
        t2.rows[i].cells[1].text = v
        t2.rows[i].cells[0].width = Inches(1.1)
        t2.rows[i].cells[1].width = Inches(5.2)
        if i == 0:
            set_cell_bg(t2.rows[0].cells[0], "2C3E50")
            set_cell_bg(t2.rows[0].cells[1], "2C3E50")
        for cell in t2.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 6. HASZNÁLATI UTASÍTÁS
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "6. Használati utasítás", level=1)
    add_body(doc, "Kérjük, olvassa el figyelmesen az első bekapcsolás előtt.")

    # 6.1
    add_heading(doc, "6.1 Állási idő használat előtt", level=2)
    add_body(doc,
        "A gép átvételét vagy mozgatását követően 24 óráig állni kell, mielőtt bekapcsolja. "
        "Ha ezt nem tartja be, a következményekért saját maga felelős."
    )

    add_heading(doc, "6.1.1 Állási idő átvétel után", level=3)
    add_body(doc,
        "Hosszú szállítás során elkerülhetetlenül rezgések és billentések fordulnak elő, "
        "amelyek hatására a kompresszor kenőolaja kifolyhat. Szükséges 24 óráig várni, hogy "
        "az olaj visszafolyjon, mielőtt a gépet használja. Ha az olaj a hűtési rendszerbe "
        "kerül, komolyan rontja a hűtési hatékonyságot."
    )

    add_heading(doc, "6.1.2 Állási idő mozgatás előtt", level=3)
    add_body(doc,
        "Ne döntse meg és ne szállítsa hosszú távolságra. Mozgatáskor vagy leürítéskor ne "
        "döntse meg és ne fordítsa fel. Ha megdöntötték vagy mozgatták, 24 óráig kell állnia "
        "mielőtt újra használható."
    )

    # 6.2
    add_heading(doc, "6.2 Berendezés összeszerelése", level=2)

    add_heading(doc, "6.2.1 Vízfeltöltési művelet", level=3)
    vizfeltoites = [
        "A hűtővizet havonta kell cserélni. Kizárólag tiszta (desztillált) vizet használjon - "
        "ásványvizet vagy csapvizet nem. 3 × 500 ml palack szükséges.",
        "Túlfolyó: Illessze be a rövid kiöntőt a túlfolyó nyílásba kattanásig. "
        "A kiöntő egyensúlyban tartja a tartályban lévő nyomást.",
        "Vízbeöntő: Tartsa a túlfolyóba beillesztett kiöntőt a helyén; a tölcsérrel ellátott "
        "vékony végű vízcsövet illessze be a vízbeöntő nyílásba kattanásig.",
        "Telítettségi jelzés: Töltsön be tiszta vizet tölcsér segítségével, amíg a "
        "túlfolyóból víz nem folyik ki - ekkor a tartály tele van.",
    ]
    for b in vizfeltoites:
        add_bullet(doc, b)

    add_img(doc, img("p8_img4.jpeg"), width_inch=4.0)
    add_body(doc,
        "A kép bal oldalán: Vízbemenet (Water Inlet), Túlfolyó (Water Spilt), "
        "Vízkimenet (Water Outlet). A tartály kb. 3 litert tartalmaz; legalább 2,5-3,0 "
        "liter desztillált vizet adjon hozzá. Feltöltés után a túlfolyóból kifolyó víz "
        "jelzi a teli szintet.", size=9
    )

    add_heading(doc, "6.2.2 Leürítési művelet", level=3)
    add_bullet(doc, "Csatlakoztassa a leürítőt: illessze be a rövid kiöntőt a leürítő nyílásba kattanásig.")
    add_bullet(doc, "Leürítési művelet: A vízcsö beillesztése után automatikusan elindul a vízelvezetés. Medencével fogja fel a lefolyó vizet.")
    add_bullet(doc, "Befejezés: Amikor nem folyik több víz, a tartály teljesen ki van ürítve.")

    # 3 kép egymás mellé - közelítő megoldás: egyenként, kis méretben
    p_row = doc.add_paragraph()
    p_row.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for fn in ["p9_img5.jpeg", "p9_img6.jpeg", "p9_img7.jpeg"]:
        if os.path.exists(img(fn)):
            p_row.add_run().add_picture(img(fn), width=Inches(1.7))
            p_row.add_run("  ")

    add_heading(doc, "6.2.3 Kézdarab-tartó felszerelése", level=3)
    add_body(doc,
        "A berendezés oldalán lévő ház mindhárom oldalán három csavarlyuk jelölést talál. "
        "Süllyesztett fejű csavarokkal szerelje fel a kézdarab-tartót."
    )

    add_heading(doc, "6.2.4 Kézdarab csatlakoztatása", level=3)
    add_body(doc,
        "Ujjaival nyomja meg a kézdarab csatlakozójának két pontját és nyomja határozottan "
        "a csatlakozóba, amíg kattanást nem hall a gyors-csatlakozón - ellenőrizze, hogy "
        "nincs vízszivárgás. Szétszerelésnél: nyomja meg a két pontot és húzza ki. "
        "A gyenge csatlakozás riasztást és szivattyúkárosodást okozhat."
    )

    add_heading(doc, "6.2.5 Tápcsatlakozó", level=3)
    add_body(doc,
        "A tápkábel egyik végét a főegységhez, másik végét az elosztóhoz csatlakoztassa. "
        "Első alkalommal távolítsa el a képernyő védőfóliáját - ez érzékenyebbé teszi az érintést."
    )

    doc.add_page_break()

    # 6.3
    add_header_line(doc)
    add_heading(doc, "6.3 Szőrtelenítés előkészítése", level=2)

    add_heading(doc, "6.3.1 Fotó készítése", level=3)
    add_body(doc,
        "Minden szőrtelenítési vagy szépségipari kezelés előtt (borotválkozás előtt is) "
        "jó megvilágításban készítsen fényképet az összehasonlíthatóság érdekében."
    )

    add_heading(doc, "6.3.2 Terület borotválása", level=3)
    add_bullet(doc, "Az ügyfélnek egy hétig kerülnie kell a napsugárzást és más bőrkárosodást.")
    add_bullet(doc, "Ha a szőrtelenítési területen smink vagy bőrápolószer van, azt előbb le kell távolítani.")
    add_bullet(doc, "Utasítsa az ügyfelet a szőr óvatos borotválására - ne karcoljuk a bőrt; "
                    "ne használjunk gyantát, epilátort, szőrtelenítő krémet.")

    add_heading(doc, "6.3.3 Bőrhűtés", level=3)
    add_body(doc,
        "Vigyen fel vastag réteg hideg gélt az ügyfél szőrtelenítési területére. Érzékeny "
        "területek kezelésekor a fájdalom csökkentése érdekében kezelés előtt jégcsomagot "
        "is alkalmazhat."
    )

    # 6.4
    add_heading(doc, "6.4 Lézer szőrtelenítési kezelés", level=2)

    add_heading(doc, "6.4.1 Bekapcsolás", level=3)
    add_body(doc, "(1) Fordítsa a kulcsot jobbra a bekapcsoláshoz. Vészhelyzet esetén azonnal nyomja meg a "
                  "piros vészleállító gombot. Ha korábban megnyomta, fél fordulattal jobbra kell "
                  "visszaállítani az újraindításhoz.")
    add_body(doc, "(2) Működés közben a vészleállító gomb megnyomásával azonnali leállítás lehetséges.")
    add_body(doc, "(3) Kétnyelvű felületen válasszon nyelvet az üdvözlőképernyőn.")
    add_body(doc, "(4) Kikapcsoláshoz forgassa vissza a kulcskapcsolót. Biztonság érdekében húzza ki "
                  "a tápkábelt is.")

    add_heading(doc, "6.4.2 Kezelési mód kiválasztása", level=3)
    add_body(doc, "Három üzemmód közül választhat:")
    add_bullet(doc, "HR  (Hair Removal)       - Professzionális szőrtelenítés")
    add_bullet(doc, "FHR (Fast Hair Removal)  - Gyors szőrtelenítés")
    add_bullet(doc, "SR  (Skin Rejuvenation)  - Bőrfiatalítás")

    # UI képernyőképek
    p_row2 = doc.add_paragraph()
    p_row2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for fn in ["p10_img8.jpeg", "p10_img9.jpeg"]:
        if os.path.exists(img(fn)):
            p_row2.add_run().add_picture(img(fn), width=Inches(2.6))
            p_row2.add_run("  ")
    add_body(doc, "1. felület: Üdvözlőképernyő     2. felület: Üzemmód-választás", size=9)

    add_heading(doc, "6.4.3 Funkció kiválasztása", level=3)
    add_body(doc, "Válassza ki: (1) nem, (2) bőrszín, (3) szőrtelenítési terület. "
                  "Az ügyfél tényleges állapota alapján válasszon, majd kattintson a Megerősítés gombra.")

    p_row3 = doc.add_paragraph()
    p_row3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for fn in ["p11_img10.jpeg", "p11_img11.jpeg", "p11_img12.jpeg"]:
        if os.path.exists(img(fn)):
            p_row3.add_run().add_picture(img(fn), width=Inches(1.75))
            p_row3.add_run("  ")
    add_body(doc, "Bal: Nők beállítása (1-2. felület)     Jobb: Férfiak beállítása (1-3. felület)     "
                  "Közép: Paraméterbeállítás", size=9)

    add_heading(doc, "6.4.4 Paraméterbeállítás", level=3)
    paramok = [
        ("(1)", "Sötét, sűrű szőr és érzékeny területek esetén az energiát alacsony értékről kell "
                "kezdeni (pl. 6 J), a frekvenciát 1-ről."),
        ("(2)", "Világos, ritka szőr és érzéketlen területek esetén az energia mérsékelt értékről "
                "kezdhető (pl. 10 J), a frekvencia 2-ről vagy 3-ról."),
        ("(3)", "Az impulzusszélességet nem kell beállítani - automatikusan igazodik az energia- "
                "és frekvenciabeállításhoz."),
        ("(4)", "A bőrhűtési paramétert 80-90%-ra kell növelni."),
    ]
    for num, szov in paramok:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        r1 = p.add_run(num + "  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(szov)
        r2.font.size = Pt(10)

    add_heading(doc, "6.4.5 Kezelési művelet", level=3)
    add_body(doc,
        'A beállítások elvégzése után kattintson a "Kész" (READY) gombra. Az előkészítési '
        "idő legfeljebb 30 másodperc; ha tovább tart, a rendszer automatikusan készenléti "
        "állapotba vált."
    )
    muveletek = [
        ("(1)", "Kezelés közben folyamatosan igazítsa az energiát és a frekvenciát. Javasolt "
                "az ügyfél által elviselhető maximális energiasűrűséggel csúsztató technikával "
                "dolgozni. A kezelőfej teljes felületen érintkezzen a bőrrel. Optimális érzés: "
                "enyhe szúrás vagy hangyacsípés."),
        ("(2)", "Jegyezze fel a beállított lézerparamétereket - a következő kezelésnél "
                "ugyanezeket lehet alkalmazni."),
        ("(3)", "Minél magasabb az energia és a frekvencia, annál jobb az eredmény. Az ügyfél "
                "által elviselhető enyhe szúrás az optimális - túl magas energia égési sérülést "
                "okozhat."),
        ("(4)", "Általánosságban csúsztató módszerrel dolgozzon; kis vagy nem csúsztatható "
                "területeken pontszerű kezelés is alkalmazható."),
    ]
    for num, szov in muveletek:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        r1 = p.add_run(num + "  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(szov)
        r2.font.size = Pt(10)

    doc.add_page_break()
    add_header_line(doc)

    add_heading(doc, "6.4.6 Gyors szőrtelenítés (FHR)", level=3)
    add_body(doc,
        "(1) Csak a bőrhűtést kell beállítani (80-90%-ra). A többi paramétert a gép előre beállítja."
    )
    add_body(doc,
        "(2) Minden testterületnél a legalacsonyabb energiaszintről kell kezdeni."
    )
    add_img(doc, img("p12_img13.jpeg"), width_inch=3.8)
    add_body(doc, "Gyors szőrtelenítés kezelőfelülete - előre beállított energiaszintek.", size=9)

    add_heading(doc, "6.4.8 A lézer szőrtelenítés elve", level=3)
    add_body(doc,
        "(1) A lézer szőrtelenítés a melanin fényelnyelő tulajdonságát használja fel: "
        "a melanin absorbeálja a lézerfényt, hőhatást generál, és elpusztítja a szőrtüszőt. "
        "Mivel a szőrtüsző pusztulása visszafordíthatatlan, a lézer tartós szőrtelenítést biztosít."
    )
    add_body(doc,
        "(2) A szőrnek növekedési ciklusa van: növekedés - hanyatlás - nyugalmi szakasz. "
        "A hanyatlás és nyugalmi szakaszban lévő szőrtüszők nem tartalmaznak melanint, ezért "
        "a lézer nem tudja elpusztítani őket. A szőr természetes növekedési ciklusa átlagosan "
        "6-8 hét. Ezért javasolt 21-30 napos kezelési időköz 5-8 alkalmon keresztül, több mint "
        "fél éven át - így az összes szőrtüsző elpusztítható. Ha fél év után marad némi finom "
        "szőr, 2-3 megerősítő kezeléssel elérhető a teljes tartós szőrtelenítés."
    )

    add_heading(doc, "6.4.9 Kezelés utáni gondozás", level=3)
    utokez = [
        "(1) Szőrtelenítés után jégcsomagot vagy aloe vera gélt alkalmazhat.",
        "(2) Ne változtasson a napi bőrápolási rutinján; kerülje az irritáló kozmetikumokat.",
        "(3) Ha a bőrben implantátum van, vagy mikroinjekciós kezelést végeztek, a kezelés nem ajánlott.",
        "(4) Néhány ügyfélen kiütés jelentkezhet - ez normális jelenség, 3-7 nap alatt elmúlik.",
        "(5) A kezelt területet 1 napig kerüljük a forró vízzel, 1 hétig a napsugárzással; nyáron "
            "fényvédőt kell használni.",
        "(6) Alkohol fogyasztása és csípős, irritáló ételek kerülendők.",
        "(7) Gyantás szőrtelenítés után lézeres kezelés csak 6 hónap után alkalmazható.",
        "(8) Az 1-2. kezelés után a szőr kihúzható lehet - de ne tegye! A szőr nélküli szőrtüsző "
            "nem vezet hőt, ami rontja a szőrtelenítés hatékonyságát.",
    ]
    for s in utokez:
        add_body(doc, s, size=10)

    # 6.5
    add_heading(doc, "6.5 Lézer bőrfiatalítási kezelés", level=2)
    add_img(doc, img("p13_img14.jpeg"), width_inch=3.8)
    add_body(doc, "Bőrfiatalítás kezelőfelülete (SOFT / COMFORTABLE / EFFECTIVE szintek).", size=9)

    add_heading(doc, "6.5.1 Kezelési módszer", level=3)
    add_body(doc,
        "(1) Csak a bőrhűtést kell beállítani (80-90%-ra).\n"
        '(2) A többi paramétert a gép előre beállítja. Új ügyfelek esetén a "Lágy fehérítés" '
        "(SOFT) üzemmóddal kezdjen."
    )

    add_heading(doc, "6.5.2 Elv", level=3)
    add_body(doc,
        "Az Oriental-Laser fő alkalmazási hullámhossza 808-810 nm. Az ilyen közel-infravörös "
        "fényt elnyelő kromofórok közé tartozik a melanin, hemoglobin és a víz. A lézer ezek "
        "elnyelése révén hőhatást vált ki:"
    )
    add_body(doc,
        "(1) A félvezető lézer a kollagén sejtekre és kötőszövetre hat: a dermiszben új "
        "kollagén sejtek és extracelluláris mátrix szintézisét indukálja, ráncokat, hegesedéseket "
        "és rugalmasság-csökkenést javít anélkül, hogy az epidermiszt károsítaná."
    )
    add_body(doc,
        "(2) A lézer melanint és kapillárisokat céloz meg termikus bontás céljából. Csökkenti "
        "a pigmentációt, mérsékeli a telangiektáziát, javítja a bőr simaságát és fényét. "
        "Kezelhető problémák: fotóöregedés, öregségi foltok, érpókháló, szürke bőr, rosacea "
        "és egyéb vaszkuláris és pigmentációs rendellenességek."
    )

    add_heading(doc, "6.5.3 Utókezelés és eredmények", level=3)
    add_body(doc,
        "(1) Egy teljes arckezelés ideje 10-20 perc (enyhe állapot: 10 perc, súlyosabb: "
        "20 perc). Két kezelés között általában 15 nap az intervallum."
    )
    add_body(doc,
        "(2) A betegek több mint 90%-ánál szignifikáns javulás tapasztalható (telangiektázia, "
        "rendszertelen pigmentáció, ráncok, érdes bőr, kitágult pórusok)."
    )
    add_body(doc,
        "(3) Nincs leállási idő. Ritkán ödéma vagy átmeneti pigmentációs változás fordulhat "
        "elő, ami magától visszatér."
    )
    add_body(doc, "(4) A kezelés után új kollagén szintézis mutatható ki a bőrben.")

    doc.add_page_break()

    # ══════════════════════════════════════
    # 7. KEZELŐFELÜLET PARAMÉTEREI
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "7. Kezelőfelület paramétereinek meghatározása", level=1)

    add_img(doc, img("p14_img15.jpeg"), width_inch=4.5)

    kepernyo_params = [
        ("7.1 Fluence (Energiasűrűség)",
         "Az egységnyi területre eső egyetlen lézerimpulzus energiaértéke. "
         "1-120 J/cm² között állítható."),
        ("7.2 Frekvencia",
         "Másodpercenkénti lézerimpulzusok száma. 1 Hz-10 Hz között állítható 1 Hz-es "
         "lépésközzel. A kiválasztott energiasűrűséghez a rendszer automatikusan az elérhető "
         "maximális frekvenciát állítja be. A '+' és '−' gombokkal kézzel is állítható."),
        ("7.3 Részösszeg",
         "Az LCD képernyő jobb felső sarkában a kezelés teljes lézerenergiáját mutatja. "
         "A 'Törlés' gombbal nullázható, vagy a gép kikapcsolásakor automatikusan törlődik."),
        ("7.4 Bőrhűtési szint",
         "A kezelőfej zafír ablakának és alumínium végének hűtési hatékonyságát jelzi. "
         "0-5 °C-os bőrfelületi hőmérsékletet tart fenn. '+' és '−' gombokkal állítható. "
         "Minél hosszabb a zöld csík, annál erősebb a hűtés. Bekapcsoláskor ki van kapcsolva."),
        ("7.5 Állapot (STATUS)",
         "A gép aktuális működési állapotát mutatja."),
    ]
    for cim, szov in kepernyo_params:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        bold_run(p, cim + ": ", size=10)
        r = p.add_run(szov)
        r.font.size = Pt(10)

    statusz_tabla = [
        ("Állapot", "Leírás"),
        ("STANDBY (Készenlét)",
         "A kezelő bármikor állíthatja a bőrhűtési szintet. 1-2 percet kell várni a "
         "lézerfej lehűlésére."),
        ("READY (Kész)",
         "A 'READY' gomb megnyomása után. A lézer készen áll; a kézdarab gombja vagy "
         "a lábkapcsoló aktiválja."),
        ("WORKING (Működik)",
         "A kézdarab gombja vagy a lábkapcsoló aktív. A lézer lő."),
        ("Riasztás",
         "Az állapotsor meghibásodás esetén riasztásjelzéseket is megjelenít. "
         "Részletek a 10. fejezetben."),
    ]
    ts = doc.add_table(rows=len(statusz_tabla), cols=2)
    ts.style = "Table Grid"
    for i, (k, v) in enumerate(statusz_tabla):
        ts.rows[i].cells[0].text = k
        ts.rows[i].cells[1].text = v
        ts.rows[i].cells[0].width = Inches(1.8)
        ts.rows[i].cells[1].width = Inches(4.5)
        if i == 0:
            set_cell_bg(ts.rows[0].cells[0], "2C3E50")
            set_cell_bg(ts.rows[0].cells[1], "2C3E50")
        for cell in ts.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # Statusbar képernyőkép
    p_sb = doc.add_paragraph()
    p_sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for fn in ["p15_img16.jpeg", "p15_img17.jpeg"]:
        if os.path.exists(img(fn)):
            p_sb.add_run().add_picture(img(fn), width=Inches(2.2))
            p_sb.add_run("   ")

    add_body(doc, "Képernyő alján megjelenő állapotsor:", size=10)
    alsok = [
        ("Vízhőmérséklet (WATER TEMP)",   "A hűtővíz hőmérséklete a készüléken belül."),
        ("Áramlás (FLOW RATE)",           "A hűtővíz áramlási sebessége L/perc-ben."),
        ("Ionitás (CONDUCTIVITY)",        "A hűtővíz vezetőképessége és tisztasága (μS/cm)."),
        ("Összes lövés (TOTAL)",          "A kezelőfej kumulatív lövésszáma (×1000 lövés)."),
    ]
    two_col_table(doc, alsok)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 8. KARBANTARTÁS
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "8. Karbantartás", level=1)
    karbantartas = [
        "8.1  Kizárólag tiszta (desztillált) vizet töltsön be. Ellenkező esetben a készülék hamar megrongálódhat.",
        "8.2  A gépet légkondicionált helyiségben kell üzemeltetni; a szobahőmérséklet 5-28 °C között legyen.",
        "1.   A hőmérséklet nem csökkenhet nulla alá - jég képződhet, ami megrongálhatja a belső "
             "víztartályokat, vízcsatornákat és a kézdarabot.",
        "2.   Tisztítsa meg a gélt a kézdarabról, és ügyeljen arra, hogy a zafír körül ne legyen por.",
        "3.   Ne hajlítsa meg erősen a kézdarab csövét - ez alacsony vízáramlási figyelmeztetést okozhat.",
        "4.   Ne csatlakoztassa és válassza le a kézdarab csatlakozóját túl gyakran - "
             "ez vízszivárgást, víztorlódást vagy rossz elektromos kapcsolatot okozhat.",
        "5.   Ne süsse el véletlenszerűen a lézert - égési sérülést vagy tüzet okozhat.",
        "6.   Mindig tartsa a kézdarabot a főegységen lévő tartón, ha nem használja. Ellenkező "
             "esetben a kondenzált víz visszafolyhat a lézerbe, ami lézerkárosodást okozhat.",
    ]
    for s in karbantartas:
        add_body(doc, s, size=10)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 9. GARANCIÁLIS FELTÉTELEK
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "9. Garanciális feltételek", level=1)
    add_body(doc,
        "Kérjük, a fenti figyelmeztetések és előírások betartásával üzemeltesse a berendezést. "
        "Az alábbi esetekben a garancia NEM érvényes:"
    )
    garancia = [
        "Víz nélküli üzemeltetés.",
        "Bekapcsolás a kézdarab csatlakoztatása nélkül.",
        "A kézdarab rendszeres leejtése.",
        "Helytelen bemeneti hálózati feszültség.",
        "Egyéb, nem rendeltetésszerű kezelési hibák.",
    ]
    for i, s in enumerate(garancia, 1):
        add_body(doc, f"{i}. {s}", size=10)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 10. HIBAELHÁRÍTÁS
    # ══════════════════════════════════════
    add_header_line(doc)
    add_heading(doc, "10. Egyszerű hibaelhárítás", level=1)
    add_body(doc,
        "A bekapcsolás után a rendszer folyamatosan önellenőrző programot futtat. "
        "Ha bármely összetevőben rendellenesség tapasztalható, riasztás jelenik meg."
    )

    # 10.1
    add_heading(doc, "10.1 Az indítási üdvözlőképernyőn megjelenő riasztás", level=2)
    t101 = doc.add_table(rows=2, cols=3)
    t101.style = "Table Grid"
    for j, h in enumerate(["Riasztás képe", "Értelmezés", "Megoldás"]):
        t101.rows[0].cells[j].text = h
        set_cell_bg(t101.rows[0].cells[j], "2C3E50")
        for para in t101.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Kép cella
    p_cell = t101.rows[1].cells[0].paragraphs[0]
    if os.path.exists(img("p17_img18.jpeg")):
        p_cell.add_run().add_picture(img("p17_img18.jpeg"), width=Inches(1.4))
    t101.rows[1].cells[1].text = (
        "A lézerkézdarab nincs megfelelően beillesztve."
    )
    t101.rows[1].cells[2].text = (
        "1. Ellenőrizze a kézdarab csatlakozását - gyenge kontaktus.\n"
        "2. Ha a dugaszon belüli kis áramköri lap megrongálódott, vegye fel a kapcsolatot "
        "a gyártóval cserére.\n"
        "Megjegyzés: A program frissítésével a riasztás maszkolható (vészhelyzeti megoldás, "
        "nem ajánlott)."
    )
    for i in range(1, 3):
        for para in t101.rows[1].cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    # 10.2
    add_heading(doc, "10.2 A beállítási felületen megjelenő riasztások", level=2)

    riasztasok = [
        {
            "img": "p17_img19.jpeg",
            "ertelm": "Állapot: Magas vízhőmérséklet - Túlhőmérséklet-riasztás.\n"
                      "A hűtővíz hőmérsékletének 15-28 °C között kell lennie. "
                      "Ha meghaladja a 34 °C-ot, a lézer leáll.",
            "megoldas": (
                "1. Ha a kompresszor normálisan működik, de a hőmérséklet magas: csökkentse "
                "a szobahőmérsékletet 30 °C alá, és várjon kb. fél órát.\n\n"
                "2. Ellenőrizze a kompresszor működését (érintse meg - érez-e rezgést, hall-e hangot?). "
                "Ha nem indul be, vegye fel a kapcsolatot a gyárral.\n\n"
                "3. Ellenőrizze a víztartály töltöttségét. Ha kevés a víz, a kompresszor relé kibillenhet.\n\n"
                "Ha elegendő a víz és a kompresszor rendesen működik, de a hőmérő 34 °C alatt "
                "mutat, a hőérzékelő vagy a vezérlőlap megrongálódhatott:\n"
                "  a) Ellenőrizze a szonda és az áramkör közötti kapcsolatot.\n"
                "  b) Cserélje ki a vezérlőlapot (lásd a mellékletet).\n"
                "  c) Ha a csere sem oldja meg, cserélje ki a hőszondát."
            ),
        },
        {
            "img": "p18_img20.jpeg",
            "ertelm": "Állapot: Alacsony vízáramlás.\n"
                      "A keringő hűtővíz sebessége alacsonyabb, mint 1,6 L/perc.",
            "megoldas": (
                "Ha az interfész '0 L/perc' értéket mutat:\n"
                "  1) A készülék kézdarab nélkül indult el - illessze be.\n"
                "  2) Laza kapcsolat az áramlásérzékelő és a vezérlőlap között.\n"
                "  3) A vezérlőlap megrongálódott (általában az OW1 optocoupler).\n"
                "  4) Az áramlásérzékelő megrongálódott.\n\n"
                "Ha 1,6 L/perc alatt mutat:\n"
                "  1) A kézdarab nincs megfelelően érintkezve a csatlakozóval - csatlakoztassa újra.\n"
                "  2) Légnyomás-egyensúlyhiány - nyissa ki a vízbeöntő csatlakozót.\n"
                "  3) Elöregedett szivattyú - cserélje ki. Ha szükséges, cserélje az ionszűrőt is "
                "(ha sötét, el van tömődve). Beépítésnél figyeljen a nyílirányra!"
            ),
        },
        {
            "img": "p19_img21.jpeg",
            "ertelm": "Állapot: Kérjük, cserélje ki a hűtővizet és az ionszűrőt.",
            "megoldas": (
                "1. Ha az ionmérő 20 μS/cm felett mutat: cserélje le a vizet tiszta vízre, "
                "és várjon fél óráig, amíg az érték 20 alá csökken.\n\n"
                "2. Ha az ionmérő 20 alatt mutat, de a riasztás megmarad:\n"
                "  a) Ellenőrizze a vezérlőlap és az ionmérő csatlakozását.\n"
                "  b) Frissítse a programot (vészhelyzeti megoldás, nem ajánlott)."
            ),
        },
        {
            "img": "p19_img22.jpeg",
            "ertelm": "Állapot: 2-es hiba (ERROR 02)",
            "megoldas": (
                "1. Ellenőrizze, hogy a vezérlőlap hűtőborda ventilátora megfelelően működik-e.\n"
                "2. Gyenge hőelvezetés a gép belsejében, vagy túl magas a környezeti hőmérséklet. "
                "Lehűlés után indítsa újra."
            ),
        },
    ]

    t102 = doc.add_table(rows=1 + len(riasztasok), cols=3)
    t102.style = "Table Grid"
    for j, h in enumerate(["Riasztás képe", "Értelmezés", "Megoldás"]):
        t102.rows[0].cells[j].text = h
        set_cell_bg(t102.rows[0].cells[j], "2C3E50")
        for para in t102.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, r in enumerate(riasztasok):
        row = t102.rows[i + 1]
        p_img = row.cells[0].paragraphs[0]
        fn = r["img"]
        if os.path.exists(img(fn)):
            p_img.add_run().add_picture(img(fn), width=Inches(1.2))
        row.cells[1].text = r["ertelm"]
        row.cells[2].text = r["megoldas"]
        for k in (1, 2):
            for para in row.cells[k].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # 10.3
    add_heading(doc, "10.3 A 'Kész' gomb megnyomása után megjelenő riasztás", level=2)

    t103 = doc.add_table(rows=2, cols=3)
    t103.style = "Table Grid"
    for j, h in enumerate(["Riasztás képe", "Értelmezés", "Megoldás"]):
        t103.rows[0].cells[j].text = h
        set_cell_bg(t103.rows[0].cells[j], "2C3E50")
        for para in t103.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p_img103 = t103.rows[1].cells[0].paragraphs[0]
    if os.path.exists(img("p20_img23.jpeg")):
        p_img103.add_run().add_picture(img("p20_img23.jpeg"), width=Inches(1.2))
    t103.rows[1].cells[1].text = (
        "A lézermodul megszakadt: Megszakadt vagy gyenge kapcsolat a lézermodul és a "
        "kézdarabban lévő kábel között."
    )
    t103.rows[1].cells[2].text = (
        "1. Ellenőrizze a kábelt a vezérlőlaptól a kézdarabban lévő lézermodulig - "
        "van-e laza csatlakozás.\n"
        "2. Ha a lézermodul belső megszakadása áll fenn, cserélje ki. "
        "Vegye fel a kapcsolatot a gyártóval."
    )
    for k in range(3):
        for para in t103.rows[1].cells[k].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    # 10.4
    add_heading(doc, "10.4 A kézdarab gombjának / lábkapcsolónak megnyomásakor megjelenő riasztások", level=2)

    t104 = doc.add_table(rows=3, cols=3)
    t104.style = "Table Grid"
    for j, h in enumerate(["Riasztás képe", "Értelmezés", "Megoldás"]):
        t104.rows[0].cells[j].text = h
        set_cell_bg(t104.rows[0].cells[j], "2C3E50")
        for para in t104.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 1. sor
    p_img104a = t104.rows[1].cells[0].paragraphs[0]
    if os.path.exists(img("p20_img24.jpeg")):
        p_img104a.add_run().add_picture(img("p20_img24.jpeg"), width=Inches(1.2))
    t104.rows[1].cells[1].text = (
        "Állapot: Helytelen működés.\n"
        "A lézer aktiválása előtt 'Kész' állapotban kell lenni. Más állapotban a gomb "
        "megnyomásakor ez a hibaüzenet jelenik meg."
    )
    t104.rows[1].cells[2].text = (
        "Engedje fel a kapcsolót, és nyomja meg a képernyőn a 'READY' (Kész) gombot."
    )

    # 2. sor
    p_img104b = t104.rows[2].cells[0].paragraphs[0]
    if os.path.exists(img("p20_img25.jpeg")):
        p_img104b.add_run().add_picture(img("p20_img25.jpeg"), width=Inches(1.2))
    t104.rows[2].cells[1].text = (
        "Állapot: Kérjük, ellenőrizze a dióda lézert vagy a dióda meghajtót.\n"
        "A rendszer azt észlelte, hogy kettőnél több lézerrúd megrongálódott, vagy a "
        "kapcsolóüzemű tápegység kimeneti feszültsége túl magas."
    )
    t104.rows[2].cells[2].text = (
        "Először ellenőrizze a lézerrudakat:\n"
        "1. Kapcsolja ki, indítsa újra, lépjen 'Kész' állapotba. Nézze meg hány lézerrúd "
        "világít halványan pirossal, és hány nem világít.\n\n"
        "2. Ha 3 vagy több lézerrúd megrongálódott, vegye fel a kapcsolatot a gyárral.\n\n"
        "3. Ha csak 1-2 nem világít: frissítse az alaplaprogramot, vagy csökkentse a tápegység "
        "kimeneti feszültségét 2 V-tal minden megrongálódott rúd után.\n"
        "(Ez vészhelyzeti megoldás - hosszú távon új lézer cseréje ajánlott.)"
    )
    for i in (1, 2):
        for k in range(3):
            for para in t104.rows[i].cells[k].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    add_body(doc, "- A kézikönyv vége -", size=10)

    # ══════════════════════════════════════
    # MENTÉS
    # ══════════════════════════════════════
    doc.save(OUT_PATH)
    print(f"\nKész! Mentve: {OUT_PATH}")


if __name__ == "__main__":
    build()
